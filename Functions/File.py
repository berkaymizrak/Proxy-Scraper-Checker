# File Functions
# Berkay MIZRAK
# www.BerkayMizrak.com
# www.DaktiNetwork.com


try:
    import time

    import os
    import sys

    # import numpy as np
    import pickle
    import xlrd
    import xlsxwriter

    import requests
    import json
    import ast

    from docx.shared import Inches
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

    from Functions import Progress
    from Functions import String
except Exception as e:
    print()
    print(e)
    while True:
        input('\n! ! ERROR --> A module is not installed...')


# When you compile your PyQt5 code with pyinstaller,
# exe file can throw error sometimes.
# To avoid that, I run below 2 functions before everything in PyQt5 code.
def _append_run_path():
    if getattr(sys, 'frozen', False):
        pathlist = []
        pathlist.append(sys._MEIPASS)

        _main_app_path = os.path.dirname(sys.executable)
        pathlist.append(_main_app_path)
        os.environ["PATH"] += os.pathsep + os.pathsep.join(pathlist)

def source_path(add_path="chromedriver.exe"):
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, add_path)
# USAGE -------------------
"""
_append_run_path()
driver = source_path("chromedriver.exe")
env_file = source_path(".env")
"""
# USAGE -------------------


def create_folder(folder_name, path='./', exit_all=True):
    folder_name = windows_folder_name(folder_name)
    path = path + folder_name
    try:
        if not os.path.exists(path):
            os.mkdir(path)
    except Exception as e:
        message = '--> An error occurred while creating folder. Please try again with running program as administrator or create folder by yourself.\n' \
                  'Folder Name: %s' % folder_name
        Progress.exit_app(message=message, e=e, exit_all=exit_all)

def windows_folder_name(name):
    forbidden_character_list = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    for character in forbidden_character_list:
        name = name.replace(character, '')
    return name

def save_dict_with_pprint_pformat(file, dict_as_string, exit_all=False):
    if isinstance(dict_as_string, str):
        try:
            file_name, file_extension = os.path.splitext(file)
            if file_extension != '.txt':
                file = file_name + '.txt'

            dict_as_string_list = dict_as_string.split('\n')
            save_records_list(file, dict_as_string_list, overwrite=True, exit_all=exit_all)
        except Exception as e:
            message = '--> An error occurred while creating file. --> "%s"' % file
            Progress.exit_app(message=message, e=e, exit_all=exit_all)
# USAGE -------------------
"""
import pprint
save = pprint.pformat(response)
File.save_dict_with_pprint_pformat(file=txt_file, dict_as_string=save)
"""
# USAGE -------------------

def json_dump(url=None, dictionary=None, json_file='Json_Data.json', header=None):
    if not header:
        user_agent = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.108 Safari/537.36'
        header = {"User-Agent": user_agent}

    if url:
        response = requests.get(
            url,
            headers=header,
        ).json()
    elif dictionary:
        response = dictionary
    else:
        return

    with open(json_file, 'w') as outfile:
        json.dump(response, outfile, indent=4, sort_keys=True)

def dump_data(file, data, exit_all=False):
    try:
        if isinstance(data, dict) or isinstance(data, list):
            file_name, file_extension = os.path.splitext(file)
            if file_extension != '.pickle':
                file = file_name + '.pickle'

            pickle_out = open(file, "wb")
            pickle.dump(data, pickle_out)
            pickle_out.close()
        # elif isinstance(data, list):
        #     file_name, file_extension = os.path.splitext(file)
        #     if file_extension != '.npy':
        #         file = file_name + '.npy'
        #
        #     np.save(file, data)
        else:
            message = "--> Data type is not acceptable. Data type only can be a 'list' or 'dict'."
            Progress.exit_app(message=message, exit_all=exit_all)
    except Exception as e:
        message = '--> An error occurred while creating file. --> "%s"' % file
        Progress.exit_app(message=message, e=e, exit_all=exit_all)

def read_dumped_data(file, data_type=dict, file_not_found_error=False, exit_all=True):
    result = None
    try:
        if data_type == dict:
            file_name, file_extension = os.path.splitext(file)
            if file_extension != '.pickle':
                file = file_name + '.pickle'

            if not os.path.exists(file):
                if file_not_found_error:
                    message = "! ! File couldn't be found in folder. --> '%s'" % (file)
                    Progress.exit_app(message=message, exit_all=exit_all)
                return None

            pickle_in = open(file, "rb")
            result = pickle.load(pickle_in)
        elif data_type == list:
            file_name, file_extension = os.path.splitext(file)
            if file_extension != '.npy':
                file = file_name + '.npy'

            if not os.path.exists(file):
                if file_not_found_error:
                    message = "! ! File couldn't be found in folder. --> '%s'" % (file)
                    Progress.exit_app(message=message, exit_all=exit_all)
                return None

            result = np.load(file).tolist()
        else:
            message = "--> File extension is not acceptable. File extension only can be '.pickle' or '.npy'."
            Progress.exit_app(message=message, exit_all=exit_all)
    except Exception as e:
        message = '--> An error occurred while reading file. --> "%s"' % file
        Progress.exit_app(e=e, message=message, exit_all=exit_all)

    return result

# LOGGING
"""
import logging
logging.basicConfig(
    filename=file,
    filemode='a',
    format='[%(asctime)s.%(msecs)d] - %(name)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S',
    level=logging.DEBUG
)
logging.info("Running Urban Planning") #  ROOT

logger = logging.getLogger('urbanGUI') #  Specify USER
logging.info("my log under urbanGUI user.")
"""
# LOGGING

def save_records_list(txt_file, records_list, overwrite=False, exit_all=True):
    try:
        if overwrite:
            file = open(txt_file, 'w', encoding='utf-8')
        else:
            file = open(txt_file, 'a', encoding='utf-8')

        for record in records_list:
            file.write('%s\n' % record)
        file.close()
    except Exception as e:
        message = '--> An error occurred while creating file. --> "%s"' % txt_file
        Progress.exit_app(message=message, e=e, exit_all=exit_all)
# USAGE for saving last location of automation app:
"""
File.save_records_list(txt_file=txt_file, records_list=[last_id], overwrite=True, exit_all=False)
"""
# USAGE for saving last location of automation app --------------------

def read_records_to_list(txt_file, file_not_found_error=False, exit_all=True):
    records = list()
    message = ''
    error_text = ''
    error_def = False
    try:
        file = open(txt_file, 'r', encoding='latin-1')
        for line in file:
            line = line.replace('\n', '')
            records.append(line)
        file.close()
    except FileNotFoundError:
        if file_not_found_error:
            error_def = True
            message = "--> File coulnd't be found in folder. --> '%s'" % txt_file
            error_text = None
    except Exception as e:
        error_def = True
        message = '--> An error occurred while reading file. --> "%s"' % txt_file
        error_text = e

    if error_def:
        Progress.exit_app(message=message, e=error_text, exit_all=exit_all)

    return records
# USAGE for reading last location of automation app:
"""
id_last = File.read_records_to_list(txt_file, file_not_found_error=False, exit_all=False)
if len(id_son) >= 1:
    id_last = int(id_last[0])
else:
    id_last = 0
"""
# USAGE for reading last location of automation app --------------------

def save_records_data(txt_file, val_list, message='File updating...', exit_all=True):
    # This def is for saving data with columns like excel but into the txt file
    try:
        print(message)
        file = open(txt_file, 'a', encoding='utf-8')
        for val in val_list:
            file.write(str(val))
            file.write('\n')
            file.write('-' * 20)
            file.write('\n')
        file.write('-' * 40)
        file.write('\n')

        file.close()
        print('File saved.')
    except Exception as e:
        message = '--> An error occurred while saving file.'
        Progress.exit_app(message=message, e=e, exit_all=exit_all)
# File.save_records_data(txt_file=txt_file, val_list=my_list, exit_all=False)

def read_records_data_to_dict(txt_file, show_progress=True, file_not_found_error=True, exit_all=True):
    # This def is for reading data with columns like excel but from plain text file
    read_dict = dict()
    try:
        total = 0
        file_exist = True
        try:
            file = open(txt_file, 'r', encoding='utf-8')
            for line in file:
                total += 1
            file.close()
        except FileNotFoundError:
            file_exist = False
            if file_not_found_error:
                message = "--> File coulnd't be found in folder. --> '%s'" % txt_file
                Progress.exit_app(message=message, exit_all=exit_all)

        if file_exist:
            count = 0
            file = open(txt_file, 'r', encoding='utf-8')
            now = time.time()
            time.sleep(0.01)
            new_line = True
            key = 1
            for line in file:
                if key not in read_dict.keys():
                    read_dict[key] = list()

                line = line[:-1]
                if show_progress:
                    count += 1
                    Progress.progress(
                        count=count,
                        total=total,
                        now=now,
                        message='Reading records...',
                    )

                if line == '-' * 40:
                    key += 1
                    new_line = True
                    continue
                if line == '-' * 20:
                    new_line = True
                    continue

                if new_line:
                    read_dict[key].append(line)
                else:
                    read_dict[key][-1] = read_dict[key][-1] + line
                new_line = False

            if show_progress:
                print()

            file.close()
            for key in list(read_dict.keys()):
                if not len(read_dict[key]):
                    del read_dict[key]
    except Exception as e:
        if show_progress:
            print()
        message = "--> An error occurred while reading file -> '%s'" % txt_file
        Progress.exit_app(e=e, message=message, exit_all=exit_all)

    return read_dict
# read_dict = File.read_records_data_to_dict(txt_file, show_progress=True, file_not_found_error=False, exit_all=False)

def write_ok_and_false_proxy(record_ip, error_file='Recorded FALSE Proxies.txt', ok_file='Recorded OK Proxies.txt',):
    # This saves the proxy to the file.

    error_ip_list = read_records_to_list(error_file, file_not_found_error=False, exit_all=False)
    ok_ip_list = read_records_to_list(ok_file, file_not_found_error=False, exit_all=False)

    if record_ip in ok_ip_list:
        ok_ip_list.remove(record_ip)
        save_records_list(ok_file, ok_ip_list, overwrite=True, exit_all=False)

    if record_ip not in error_ip_list:
        save_records_list(error_file, [record_ip], overwrite=False, exit_all=False)

def find_file(file, path='.'):
    if isinstance(file, str):
        if os.path.exists(file):
            folder = os.listdir(path)
            for folder_file in folder:
                if file.lower() == folder_file.lower():
                    file = folder_file
                    break
                elif String.lower_string(file) == String.lower_string(folder_file):
                    # Lower case file name with Turkish characters.
                    file = folder_file
                    break
    return file


def excel_read_to_dict(excel, number_of_sheet=0, exit_all=False):
    all_data = dict()
    headers = dict()

    try:
        # Check and add xlsx or xls if there is not at the end.
        file_name, file_extension = os.path.splitext(excel)
        if file_extension != '.xlsx' or file_extension != '.xls':
            excel = file_name + '.xlsx'

        # check all versions of the file name if it is exist in directory.
        # (Checking with all lower and capital characters for excel name if it is equal any file.)
        excel = find_file(excel)
        if not os.path.exists(excel):
            # So given file name could not be found in directory with any combinations of capital and lower characters.
            excel2 = None

            # switch between xlsx and xls
            if file_extension != '.xlsx':
                excel2 = file_name + '.xls'
            elif file_extension != '.xls':
                excel2 = file_name + '.xlsx'

            if excel2:
                # if given file name is xlsx, it switched to xls in "excel2"
                # if given file name is xls, it switched to xlsx in "excel2"
                # and checking again...
                excel2 = find_file(excel2)
                if not os.path.exists(excel2):
                    message = "! ! File couldn't be found in folder. --> '%s' or '%s'" % (excel, excel2)
                    Progress.exit_app(message=message, exit_all=exit_all)
                    return all_data, headers
                else:
                    excel = excel2
            else:
                message = "! ! File couldn't be found in folder. --> '%s'" % (excel)
                Progress.exit_app(message=message, exit_all=exit_all)
                return all_data, headers

        workbook = xlrd.open_workbook(excel)  # sheet
        sheet = workbook.sheet_by_index(number_of_sheet)  # page

        number_of_column = sheet.ncols
        number_of_row = len(sheet.col(0))

        count = 0
        total = number_of_row
        now = time.time()
        message = 'Reading excel...'
        time.sleep(0.01)

        number_of_data = 0
        number_of_header = 0

        for y in range(number_of_row):
            key = sheet.cell_value(rowx=y, colx=0)
            try:
                key = int(key)
            except:
                pass

            # I only get integer keys which means excel rows which has integer at first cell.
            # This is for not getting header rows in my dictionary.
            # and I design my excels with ID column at first column.
            if isinstance(key, int):
                number_of_data += 1
                all_data[number_of_data] = list()
                for x in range(number_of_column):
                    val = sheet.cell_value(rowx=y, colx=x)
                    val = String.float_to_integer(val, force_number=False)
                    all_data[number_of_data].append(val)
            else:
                number_of_header += 1
                headers[number_of_header] = list()
                for x in range(number_of_column):
                    val = sheet.cell_value(rowx=y, colx=x)
                    val = String.float_to_integer(val, force_number=False)
                    headers[number_of_header].append(val)

            count += 1

            Progress.progress(
                count=count,
                total=total,
                now=now,
                message=message,
            )
    except PermissionError:
        message = "--> '%s' can't access to this file.\nIt is probably because the file is open. If this excel is open, please close it and re-run program." % excel
        Progress.exit_app(message=message, exit_all=exit_all)
    except Exception as e:
        message = "--> An error occurred while reading file... '%s'" % excel
        Progress.exit_app(e=e, message=message, exit_all=exit_all)

    print('\nNumber of item: %s' % len(all_data))
    # it returns a dictionary from 3 rows excel file as:
    # all_data = {
    #     1: ['1st Column Value', '2nd Column Value', '3rd Column Value', '4th Column Value', '5th Column Value', ],
    #     2: ['1st Column Value', '2nd Column Value', '3rd Column Value', '4th Column Value', '5th Column Value', ],
    #     3: ['1st Column Value', '2nd Column Value', '3rd Column Value', '4th Column Value', '5th Column Value', ],
    # }
    return all_data, headers

def excel_create(excel, all_data, headers=None, sizes=None, locations=None, page_name='Page1', exit_all=False, ):
    if not headers:
        headers = list()

    # Check and add xlsx or xls if there is not at the end.
    file_name, file_extension = os.path.splitext(excel)
    if file_extension != '.xlsx' or file_extension != '.xls':
        excel = file_name + '.xlsx'

    try:
        message = "'%s'  --> Creating..." % excel
        total = len(all_data)
        print(message)

        if not total:
            message = '\n--> No data.'
            print(message)
            return

        if len(all_data):
            length_max = 0
            for val in all_data.values():
                try:
                    if length_max < len(val):
                    # Find the row which has maximum length
                        if isinstance(val[-1], dict):
                            length_max = len(val) - 1
                        else:
                            length_max = len(val)
                except:
                    pass

            i = 0
            while len(headers) < length_max:
                # if maximum length of any row larger than HEADERS, add "Header %i" rest of the headers
                i += 1
                headers.append('Header %s' % i)

            for key in list(all_data.keys()):
                # if length of Headers larger than any row, add empty cell end of the row
                while len(headers) > len(all_data[key]):
                    all_data[key].append('')


        if sizes:
            while len(headers) > len(sizes):
                sizes.append(20)
        else:
            sizes = list()
            for head in headers:
                sizes.append(25)

        if locations:
            while len(headers) > len(locations):
                locations.append('left')
        else:
            locations = list()
            for head in headers:
                locations.append('left')

        attrs_loc = dict()
        for val in all_data.values():
            for elem in val:
                if isinstance(elem, dict):
                    for name, attr in elem.items():
                        if name not in attrs_loc.keys():
                            headers.append(name)
                            sizes.append(20)
                            locations.append('left')

                            attrs_loc[name] = len(headers)

        workbook = xlsxwriter.Workbook(excel)
        worksheet = workbook.add_worksheet(page_name)

        worksheet.freeze_panes(1, 0)

        cell_format_header = workbook.add_format({'border': 1})
        cell_format_header.set_pattern(1)
        cell_format_header.set_bg_color('orange')
        cell_format_header.set_align('center')
        cell_format_header.set_align('vcenter')
        cell_format_header.set_bold()

        cell_format_center_regular = workbook.add_format({'border': 1})
        cell_format_center_regular.set_align('center')
        cell_format_center_regular.set_align('vcenter')

        cell_format_regular = workbook.add_format({'border': 1})
        cell_format_regular.set_align('left')
        cell_format_regular.set_align('vcenter')

        cell_format_right_regular = workbook.add_format({'border': 1})
        cell_format_right_regular.set_align('right')
        cell_format_right_regular.set_align('vcenter')

        cell_format_copyr = workbook.add_format({'border': 1})
        cell_format_copyr.set_pattern(1)
        cell_format_copyr.set_bg_color('FABF8F')
        cell_format_copyr.set_align('center')
        cell_format_copyr.set_align('vcenter')
        cell_format_copyr.set_bold()

        row = 0
        col = 0
        set_say = 0
        worksheet.write(row, col, 'ID', cell_format_header)
        worksheet.set_column(set_say, set_say, 8)
        set_say += 1
        col += 1
        for head, size in zip(headers, sizes):
            worksheet.write(row, col, head, cell_format_header)
            worksheet.set_column(set_say, set_say, size)
            set_say += 1
            col += 1
        worksheet.write(row, col, 'Automated by BerkayMizrak.com', cell_format_copyr)
        worksheet.set_column(set_say, set_say, 34)
        row += 1

        count = 0
        now = time.time()
        time.sleep(0.01)

        id_count = 0
        for val in all_data.values():
            id_count += 1
            col = 0
            worksheet.write(row, col, id_count, cell_format_center_regular)
            for elem in val:
                if isinstance(elem, dict):
                    continue

                col += 1
                if col > len(locations):
                    go_left = True
                else:
                    go_left = False
                    if locations[col - 1] == 'center':
                        worksheet.write(row, col, elem, cell_format_center_regular)
                    elif locations[col - 1] == 'right':
                        worksheet.write(row, col, elem, cell_format_right_regular)
                    else:
                        go_left = True
                if go_left:
                    try:
                        elem = int(elem)
                        worksheet.write(row, col, elem, cell_format_center_regular)
                    except:
                        worksheet.write(row, col, elem, cell_format_regular)

            for elem in val:
                if isinstance(elem, dict):
                    for name, attr in elem.items():
                        worksheet.write(row, attrs_loc[name], attr, cell_format_regular)

            row += 1

            count += 1
            Progress.progress(count=count, total=total, now=now, )

        print()
        workbook.close()
        message = "'%s' Data Saved to Excel -->> '%s'" % (count, excel)
        print(message)
    except PermissionError:
        message = "--> '%s' can't access to this file.\nIt is probably because the file is open. If this excel is open, please close it and re-run program." % excel
        Progress.exit_app(message=message, exit_all=exit_all)
    except Exception as e:
        message = "--> An error occurred while creating file... '%s'" % excel
        Progress.exit_app(e=e, message=message, exit_all=exit_all)
# USAGE -------------------
"""
DEFINED IN 'EXCEL READ AND CREATE' REPOSITORY. 
"""
# USAGE -------------------


def create_word(word, my_rows):
    try:
        # Check and add docx if there is not at the end.
        file_name, file_extension = os.path.splitext(word)
        if file_extension != '.docx':
            word = file_name + '.xlsx'

        message = "'%s'  --> Creating..." % word
        total = len(my_rows)
        print(message)

        if not total:
            message = '\n--> No data.'
            print(message)
            return

        document = Document()

        """
        EXTRA DETAILS FOR DESIGN:     
    
        document.add_heading('Document Title', 0)
        
        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True
        
        document.add_heading('Heading, level 1', level=1)
        
        document.add_picture('monty-truth.png', width=Inches(1.25))
        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        document.add_page_break()
        """

        styles = document.styles

        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        style = styles.add_style('MyHeader1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.bold = True
        font.name = 'Arial'
        font.size = Pt(14)

        style = styles.add_style('MyHeader2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Arial'
        font.size = Pt(13)

        for row in my_rows:
            table = row.get('table', False)
            if not table:
                text = row.get('text', '')
                style_name = row.get('style', None)
                location = row.get('location', 0)
                bold = row.get('bold', False)
                italic = row.get('italic', False)
                underline = row.get('underline', False)
                size = row.get('size', 12)

                paragraph = document.add_paragraph()
                run = paragraph.add_run(text)

                if style_name:
                    style = document.styles[style_name]
                    paragraph.style = style
                else:
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline

                    font = style.font
                    font.size = Pt(size)

                paragraph.alignment = location  # 0: left, 1: center, 2: right, 3: justify
            else:
                data = row.get('data', [])
                border = row.get('border', False)

                if len(data):
                    table_obj = document.add_table(rows=0, cols=len(data[0]), )
                    if border:
                        table_obj.style = 'TableGrid'

                    for count_row, table_row in enumerate(data):
                        row_cells = table_obj.add_row().cells
                        for enum, cell in enumerate(table_row):
                            row_cells[enum].text = str(cell)

                        if count_row == 0:
                            for count_cell in range(len(table_row)):
                                row_cells[count_cell].paragraphs[0].runs[0].font.bold = True
                                row_cells[count_cell].paragraphs[0].alignment = 1

        print()
        document.save(word)

        message = "Word Created -->> '%s'" % (word)
        print(message)
    except PermissionError:
        message = "--> '%s' can't access to this file.\nIt is probably because the file is open. If this word is open, please close it and re-run program." % word
        Progress.exit_app(message=message, exit_all=False)
    except Exception as e:
        message = "--> An error occurred while creating file... '%s'" % word
        Progress.exit_app(e=e, message=message, exit_all=False)
"""
USAGE:

my_rows = [
    {'text': 'Hello World. Center, bold, italic, underline, size:15', 'style': 'Normal', 'location': 1, 'bold': True, 'italic': True, 'underline': True, 'size': 15},
    {'text': 'My Special Header. Right, bold', 'style': 'MyHeader1', 'location': 2, 'bold': True, 'italic': False, 'underline': False},
    {'text': 'Hello World. Left, normal', 'style': 'Normal', 'location': 0},
    {'text': 'Lorem ipsum. Center, italic', 'style': 'Normal', 'location': 1, 'italic': True, },
    {'text': 'Specified text. Left, underline', 'style': 'Normal', 'location': 0, 'underline': True},
]  # location = 0: left, 1: center, 2: right, 3: justify

# ---- OR : 
my_rows = []  # location = 0: left, 1: center, 2: right, 3: justify

row = {
    'text': 'Hello World. Left, normal',
    'style': 'Normal',
    'location': 0
}
my_rows.append(row)

data = []
data.append(['Product', 'Quantity', 'Price', ])
data.append(['AI Builder Capacity', '1', '$324,45', ])
row = {
    'table': True,
    'data': data,
    'border': True,
}
my_rows.append(row)

row = {
    'text': 'Lorem ipsum. Center, italic',
    'style': 'Normal', 
    'location': 1, 
    'italic': True,
}
my_rows.append(row)
# --------

create_word('test', my_rows)
"""

